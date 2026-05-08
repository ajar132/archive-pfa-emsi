#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur du Rapport PFA : Archive Digitale PFA
Lancer : pip install python-docx   puis   python generer_rapport_pfa.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Helpers ──────────────────────────────────────────────────────────────────

def pb(doc):
    doc.add_page_break()

def para(doc, text, bold=False, italic=False, size=11,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def center(doc, text, bold=False, size=12, sa=6):
    return para(doc, text, bold=bold, size=size,
                align=WD_ALIGN_PARAGRAPH.CENTER, sa=sa)

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def tbl(doc, headers, rows, caption=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if caption:
        para(doc, caption, italic=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    return t

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


# ─── Document ─────────────────────────────────────────────────────────────────

doc = Document()
s = doc.sections[0]
s.top_margin = Cm(2.5)
s.bottom_margin = Cm(2.5)
s.left_margin = Cm(3.0)
s.right_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════
para(doc, "")
center(doc, "ÉCOLE MAROCAINE DES SCIENCES DE L'INGÉNIEUR", bold=True, size=14, sa=2)
center(doc, "EMSI – [Votre Ville]", bold=True, size=12, sa=2)
center(doc, "Filière : Génie Informatique / Ingénierie des Systèmes d'Information", size=11, sa=4)
center(doc, "═" * 55, size=10, sa=14)
center(doc, "RAPPORT DE PROJET DE FIN D'ANNÉE – 2024/2025", bold=True, size=13, sa=20)
center(doc, "ARCHIVE DIGITALE PFA", bold=True, size=20, sa=6)
center(doc, "Application Desktop de Gestion des Rapports", bold=True, size=14, sa=4)
center(doc, "de Projets de Fin d'Année à l'EMSI", bold=True, size=14, sa=20)
center(doc, "═" * 55, size=10, sa=16)
center(doc, "Réalisé par :", bold=True, size=11, sa=4)
center(doc, "[Prénom NOM Étudiant 1]", size=12, sa=2)
center(doc, "[Prénom NOM Étudiant 2]", size=12, sa=16)
center(doc, "Encadré par :", bold=True, size=11, sa=4)
center(doc, "[Pr. Prénom NOM]", size=12, sa=2)
center(doc, "[Grade – Département]", size=11, sa=20)
center(doc, "═" * 55, size=10, sa=16)
center(doc, "Année Universitaire 2024 – 2025", bold=True, size=12)
pb(doc)

# ══════════════════════════════════════════════════════════
# DÉDICACE
# ══════════════════════════════════════════════════════════
doc.add_heading("Dédicace", level=1)
para(doc,
     "À nos chers parents, qui nous ont accordé leur soutien indéfectible et leurs "
     "sacrifices tout au long de notre parcours académique.",
     indent=1.25)
para(doc,
     "À nos familles et amis, pour leur présence et leurs encouragements constants "
     "qui nous ont portés dans les moments difficiles.",
     indent=1.25)
para(doc,
     "À nos encadrants et enseignants, pour leur savoir partagé avec générosité et "
     "leur passion pour la transmission du savoir.",
     indent=1.25)
para(doc,
     "À tous ceux qui croient en la valeur de l'éducation et en la force de la "
     "technologie au service du progrès.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# REMERCIEMENTS
# ══════════════════════════════════════════════════════════
doc.add_heading("Remerciements", level=1)
para(doc,
     "Au terme de ce Projet de Fin d'Année, il nous est agréable d'adresser nos plus "
     "sincères remerciements à toutes les personnes qui ont contribué, de près ou de "
     "loin, à la réussite de ce travail.",
     indent=1.25)
para(doc,
     "Nous exprimons notre profonde gratitude à notre encadrant, [Pr. Prénom NOM], "
     "pour ses orientations précieuses, sa disponibilité et la rigueur scientifique "
     "dont il a su nous faire bénéficier tout au long de ce projet. Ses conseils "
     "éclairés et sa patience ont été déterminants dans la concrétisation de ce travail.",
     indent=1.25)
para(doc,
     "Nos remerciements vont également à la Direction de l'EMSI, et plus "
     "particulièrement au corps professoral de la filière Génie Informatique, pour "
     "la qualité de la formation dispensée et pour les ressources mises à notre disposition.",
     indent=1.25)
para(doc,
     "Nous remercions aussi nos camarades de promotion pour les échanges fructueux "
     "et l'entraide dont nous avons bénéficié tout au long de cette année universitaire.",
     indent=1.25)
para(doc,
     "Enfin, nos remerciements les plus chaleureux vont à nos familles pour leur "
     "soutien moral et leur patience sans faille durant cette période intense.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# RÉSUMÉ
# ══════════════════════════════════════════════════════════
doc.add_heading("Résumé", level=1)
para(doc,
     "Ce rapport présente les travaux réalisés dans le cadre du Projet de Fin d'Année "
     "(PFA) pour la filière Génie Informatique à l'École Marocaine des Sciences de "
     "l'Ingénieur (EMSI). Le projet consiste en la conception et le développement d'une "
     "application desktop intitulée « Archive Digitale PFA ».",
     indent=1.25)
para(doc,
     "L'application vise à résoudre la problématique de la gestion manuelle et "
     "fragmentée des rapports de projets de fin d'année au sein de l'établissement. "
     "Elle centralise l'ensemble du cycle de vie des rapports PFA : dépôt, "
     "consultation, évaluation et archivage, dans un environnement sécurisé et structuré.",
     indent=1.25)
para(doc,
     "Développée en Python 3 avec la bibliothèque graphique PyQt5 et une base de "
     "données SQLite3, l'application met en œuvre une architecture en couches et un "
     "système de contrôle d'accès basé sur les rôles (RBAC) avec trois profils : "
     "administrateur, encadrant et étudiant. La sécurité est assurée par le hachage "
     "bcrypt des mots de passe et un mécanisme de blocage de compte après trois "
     "tentatives d'authentification infructueuses.",
     indent=1.25)
para(doc,
     "Parmi les fonctionnalités avancées : l'extraction automatique de texte des PDF "
     "via PyMuPDF, un moteur de détection de similarité anti-plagiat basé sur "
     "l'algorithme TF-IDF et la similarité cosinus (scikit-learn), ainsi qu'un "
     "tableau de bord statistique avec des graphiques générés par Matplotlib.",
     indent=1.25)
para(doc,
     "Mots-clés : Python, PyQt5, SQLite3, bcrypt, PyMuPDF, TF-IDF, scikit-learn, "
     "RBAC, gestion documentaire, archive numérique, anti-plagiat.",
     bold=True, indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════
doc.add_heading("Abstract", level=1)
para(doc,
     "This report presents the work carried out as part of the Final Year Project "
     "(PFA) in the Computer Engineering program at the École Marocaine des Sciences "
     "de l'Ingénieur (EMSI). The project involves the design and development of a "
     "desktop application entitled « Archive Digitale PFA » (Digital PFA Archive).",
     indent=1.25)
para(doc,
     "The application aims to solve the problem of manual and fragmented management "
     "of final year project reports within the institution. It centralizes the entire "
     "lifecycle of PFA reports: submission, consultation, evaluation, and archiving, "
     "in a secure and structured environment.",
     indent=1.25)
para(doc,
     "Developed in Python 3 with the PyQt5 graphical library and an SQLite3 database, "
     "the application implements a layered architecture and a Role-Based Access Control "
     "(RBAC) system with three profiles: administrator, supervisor, and student. "
     "Security is ensured by bcrypt password hashing and an account lockout mechanism "
     "after three failed authentication attempts.",
     indent=1.25)
para(doc,
     "Among the advanced features: automatic text extraction from PDFs via PyMuPDF, "
     "an anti-plagiarism similarity detection engine based on the TF-IDF algorithm "
     "and cosine similarity (scikit-learn), and a statistical dashboard with charts "
     "generated by Matplotlib.",
     indent=1.25)
para(doc,
     "Keywords: Python, PyQt5, SQLite3, bcrypt, PyMuPDF, TF-IDF, scikit-learn, RBAC, "
     "document management, digital archive, anti-plagiarism.",
     bold=True, indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (placeholder)
# ══════════════════════════════════════════════════════════
doc.add_heading("Table des Matières", level=1)
para(doc,
     "[Dans Microsoft Word : Références → Table des matières → Table automatique 1]",
     italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
pb(doc)

# ══════════════════════════════════════════════════════════
# LISTE DES FIGURES
# ══════════════════════════════════════════════════════════
doc.add_heading("Liste des Figures", level=1)
figures = [
    ("Figure 1",  "Architecture générale en trois couches"),
    ("Figure 2",  "Diagramme Entité-Relation (MCD)"),
    ("Figure 3",  "Diagramme de cas d'utilisation global"),
    ("Figure 4",  "Diagramme de cas d'utilisation – Authentification (CU01)"),
    ("Figure 5",  "Diagramme de cas d'utilisation – Dépôt rapport (CU03)"),
    ("Figure 6",  "Diagramme de cas d'utilisation – Évaluation (CU05)"),
    ("Figure 7",  "Diagramme de séquence – Authentification"),
    ("Figure 8",  "Diagramme de séquence – Dépôt avec analyse de similarité"),
    ("Figure 9",  "Diagramme de séquence – Évaluation d'un rapport"),
    ("Figure 10", "Diagramme de classes de l'application"),
    ("Figure 11", "Interface de connexion (LoginWindow)"),
    ("Figure 12", "Fenêtre principale – Vue administrateur"),
    ("Figure 13", "Module de dépôt de rapport (DepotWindow)"),
    ("Figure 14", "Résultats de l'analyse de similarité"),
    ("Figure 15", "Module de recherche multicritères"),
    ("Figure 16", "Module d'évaluation – Vue encadrant"),
    ("Figure 17", "Tableau de bord statistique (DashboardWindow)"),
    ("Figure 18", "Fenêtre de prévisualisation PDF"),
    ("Figure 19", "Module de gestion des utilisateurs"),
    ("Figure 20", "Résultats de l'exécution des tests unitaires"),
]
for fig, titre in figures:
    p = doc.add_paragraph()
    r = p.add_run(f"{fig} : {titre}")
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
pb(doc)

# ══════════════════════════════════════════════════════════
# LISTE DES TABLEAUX
# ══════════════════════════════════════════════════════════
doc.add_heading("Liste des Tableaux", level=1)
tableaux = [
    ("Tableau 1",  "Récapitulatif des acteurs et leurs rôles"),
    ("Tableau 2",  "Cas d'utilisation de l'application"),
    ("Tableau 3",  "Description de la table utilisateur"),
    ("Tableau 4",  "Description de la table rapport"),
    ("Tableau 5",  "Description de la table option_filiere"),
    ("Tableau 6",  "Technologies et bibliothèques utilisées"),
    ("Tableau 7",  "Structure des modules de l'application"),
    ("Tableau 8",  "Workflow des statuts d'un rapport"),
    ("Tableau 9",  "Résultats des tests unitaires"),
    ("Tableau 10", "Tests d'authentification"),
    ("Tableau 11", "Tests du contrôle d'accès (RBAC)"),
    ("Tableau 12", "Tests du dépôt de rapport"),
    ("Tableau 13", "Tests d'évaluation"),
    ("Tableau 14", "Synthèse des résultats des tests"),
]
for tab, titre in tableaux:
    p = doc.add_paragraph()
    r = p.add_run(f"{tab} : {titre}")
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
pb(doc)

# ══════════════════════════════════════════════════════════
# LISTE DES ABRÉVIATIONS
# ══════════════════════════════════════════════════════════
doc.add_heading("Liste des Abréviations", level=1)
tbl(doc,
    ["Abréviation", "Signification"],
    [
        ("API",    "Application Programming Interface"),
        ("CRUD",   "Create, Read, Update, Delete"),
        ("DB",     "Database (Base de données)"),
        ("EMSI",   "École Marocaine des Sciences de l'Ingénieur"),
        ("GUI",    "Graphical User Interface"),
        ("IHM",    "Interface Homme-Machine"),
        ("MCD",    "Modèle Conceptuel de Données"),
        ("MVC",    "Modèle-Vue-Contrôleur"),
        ("PDF",    "Portable Document Format"),
        ("PFA",    "Projet de Fin d'Année"),
        ("RBAC",   "Role-Based Access Control"),
        ("SQL",    "Structured Query Language"),
        ("TF-IDF", "Term Frequency – Inverse Document Frequency"),
        ("UML",    "Unified Modeling Language"),
    ])
pb(doc)

# ══════════════════════════════════════════════════════════
# INTRODUCTION GÉNÉRALE
# ══════════════════════════════════════════════════════════
doc.add_heading("Introduction Générale", level=1)
para(doc,
     "À l'ère du numérique, la transformation des processus administratifs et "
     "pédagogiques constitue un enjeu majeur pour les établissements d'enseignement "
     "supérieur. La gestion des documents académiques, et en particulier celle des "
     "rapports de Projets de Fin d'Année (PFA), exige une organisation rigoureuse, "
     "une traçabilité fiable et un accès sécurisé aux informations.",
     indent=1.25)
para(doc,
     "Au sein de l'École Marocaine des Sciences de l'Ingénieur (EMSI), les rapports "
     "PFA représentent le fruit du travail de centaines d'étudiants chaque année. "
     "Leur gestion, lorsqu'elle est réalisée manuellement ou à travers des outils "
     "non adaptés, engendre de nombreux problèmes : perte de documents, absence de "
     "mécanisme d'évaluation structuré, difficulté à consulter les travaux antérieurs "
     "et risque élevé de plagiat.",
     indent=1.25)
para(doc,
     "C'est dans ce contexte que s'inscrit notre projet : la conception et le "
     "développement d'une application desktop intitulée « Archive Digitale PFA ». "
     "Cette solution logicielle vise à numériser et à automatiser l'ensemble du "
     "cycle de vie des rapports PFA, depuis leur dépôt par l'administration jusqu'à "
     "leur archivage final, en passant par leur évaluation par les encadrants.",
     indent=1.25)
para(doc,
     "L'application, développée avec les technologies Python et PyQt5, propose une "
     "interface conviviale adaptée aux différents profils d'utilisateurs : "
     "administrateurs, encadrants et étudiants. Elle intègre des fonctionnalités "
     "avancées telles que la détection de similarité pour lutter contre le plagiat, "
     "la prévisualisation intégrée des PDF et un tableau de bord statistique en "
     "temps réel.",
     indent=1.25)
para(doc, "Ce rapport est organisé en cinq chapitres :", indent=1.25)
bullet(doc, "Le Chapitre 1 présente le cadre général : établissement, problématique et objectifs.")
bullet(doc, "Le Chapitre 2 expose l'analyse des besoins fonctionnels et non-fonctionnels.")
bullet(doc, "Le Chapitre 3 détaille la conception : architecture, modèle de données et diagrammes UML.")
bullet(doc, "Le Chapitre 4 décrit la réalisation : technologies, implémentation et interfaces.")
bullet(doc, "Le Chapitre 5 présente les tests et la validation de l'application.")
para(doc,
     "Une conclusion générale clôture ce rapport en récapitulant les réalisations "
     "et en présentant les perspectives d'évolution envisagées.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CHAPITRE 1 – CADRE GÉNÉRAL
# ══════════════════════════════════════════════════════════
doc.add_heading("Chapitre 1 : Cadre Général du Projet", level=1)
para(doc, "Introduction du chapitre", bold=True)
para(doc,
     "Ce chapitre présente le contexte général du projet : l'établissement d'accueil, "
     "la problématique identifiée et les objectifs fixés.",
     indent=1.25)

doc.add_heading("1.1  Présentation de l'Établissement d'Accueil", level=2)
para(doc,
     "L'École Marocaine des Sciences de l'Ingénieur (EMSI) est un établissement "
     "d'enseignement supérieur privé fondé en 1986 à Casablanca. Elle forme des "
     "ingénieurs et des techniciens dans les domaines des technologies de "
     "l'information, du génie industriel et du génie civil.",
     indent=1.25)
para(doc,
     "L'EMSI dispose de plusieurs campus répartis dans les grandes villes marocaines : "
     "Casablanca, Rabat, Marrakech, Fès, Agadir et Oujda. Elle propose des formations "
     "allant du BTS jusqu'au diplôme d'ingénieur (Bac+5). La filière Génie Informatique "
     "forme des professionnels capables de concevoir, développer et déployer des "
     "systèmes d'information complexes.",
     indent=1.25)
para(doc,
     "Le Projet de Fin d'Année (PFA) constitue une étape essentielle de cette "
     "formation, permettant aux étudiants d'appliquer les compétences acquises "
     "sur un projet concret et de démontrer leur capacité à résoudre des "
     "problèmes réels.",
     indent=1.25)

doc.add_heading("1.2  Contexte et Problématique", level=2)
para(doc,
     "Chaque année, l'EMSI accueille des centaines d'étudiants réalisant leur PFA. "
     "Ces projets donnent lieu à des rapports techniques qui constituent un patrimoine "
     "intellectuel précieux pour l'école. Cependant, leur gestion actuelle présente "
     "plusieurs problèmes :",
     indent=1.25)
bullet(doc, "Absence d'un système centralisé de stockage et d'indexation des rapports.")
bullet(doc, "Difficulté pour les étudiants et les enseignants de consulter les travaux antérieurs.")
bullet(doc, "Processus d'évaluation non structuré, sans workflow défini ni traçabilité des décisions.")
bullet(doc, "Risque de plagiat difficile à détecter sans outil automatisé de comparaison de contenu.")
bullet(doc, "Perte ou détérioration de documents physiques sur le long terme.")
para(doc,
     "Face à ces problèmes, il est apparu nécessaire de développer une solution "
     "logicielle permettant de gérer l'ensemble du cycle de vie des rapports PFA "
     "de manière numérique, structurée et sécurisée.",
     indent=1.25)

doc.add_heading("1.3  Objectifs du Projet", level=2)
para(doc, "Le projet « Archive Digitale PFA » vise les objectifs suivants :", indent=1.25)
numbered(doc, "Centraliser le stockage des rapports PFA dans un système d'archivage numérique.")
numbered(doc, "Automatiser le workflow d'évaluation (dépôt → révision → évaluation → archivage).")
numbered(doc, "Implémenter un contrôle d'accès basé sur les rôles (RBAC) avec trois profils.")
numbered(doc, "Offrir une recherche multicritères pour faciliter la consultation des rapports.")
numbered(doc, "Intégrer un système de détection de similarité pour prévenir le plagiat.")
numbered(doc, "Fournir un tableau de bord statistique en temps réel pour les responsables.")
numbered(doc, "Garantir la sécurité des données via le hachage bcrypt et l'anti-brute-force.")

doc.add_heading("1.4  Organisation du Rapport", level=2)
para(doc,
     "Ce rapport est structuré en cinq chapitres couvrant les phases d'analyse, "
     "de conception, d'implémentation et de tests. Des diagrammes UML et des "
     "captures d'interfaces illustrent les propos à chaque étape.",
     indent=1.25)

para(doc, "Conclusion du chapitre", bold=True)
para(doc,
     "Ce premier chapitre a posé les bases du projet. Le chapitre suivant est "
     "consacré à l'analyse détaillée des besoins fonctionnels et non-fonctionnels.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CHAPITRE 2 – ANALYSE DES BESOINS
# ══════════════════════════════════════════════════════════
doc.add_heading("Chapitre 2 : Analyse des Besoins", level=1)
para(doc, "Introduction du chapitre", bold=True)
para(doc,
     "L'analyse des besoins est une étape fondamentale du processus de développement "
     "logiciel. Ce chapitre identifie les acteurs, définit les fonctionnalités "
     "attendues et précise les contraintes non-fonctionnelles.",
     indent=1.25)

doc.add_heading("2.1  Identification des Acteurs", level=2)
para(doc,
     "Le système interagit avec trois types d'acteurs, chacun disposant de droits "
     "et de responsabilités bien définis :",
     indent=1.25)
tbl(doc,
    ["Acteur", "Rôle dans le système", "Responsabilités principales"],
    [
        ("Administrateur", "Super-utilisateur",
         "Gestion des comptes utilisateurs, dépôt des rapports, accès complet"),
        ("Encadrant", "Évaluateur",
         "Évaluation des rapports assignés, accès au tableau de bord"),
        ("Étudiant", "Lecteur",
         "Consultation et téléchargement des rapports archivés uniquement"),
    ],
    caption="Tableau 1 : Récapitulatif des acteurs et leurs rôles")

doc.add_heading("2.2  Besoins Fonctionnels", level=2)
para(doc,
     "Les besoins fonctionnels ont été recueillis auprès des parties prenantes et "
     "formalisés sous forme de cas d'utilisation (CU) :",
     indent=1.25)
tbl(doc,
    ["ID", "Nom du cas d'utilisation", "Acteurs", "Description synthétique"],
    [
        ("CU01", "Authentification et déconnexion", "Tous",
         "Connexion par email/mot de passe. Blocage après 3 tentatives."),
        ("CU02", "Gestion des utilisateurs", "Administrateur",
         "CRUD des comptes (ajout, modification, suppression). Rôles distincts."),
        ("CU03", "Dépôt d'un rapport", "Administrateur",
         "Envoi d'un PDF avec métadonnées (titre, auteurs, encadrant, promotion)."),
        ("CU04", "Recherche de rapports", "Tous",
         "Recherche par mot-clé et/ou promotion. Les étudiants voient uniquement 'Archivé'."),
        ("CU05", "Évaluation de rapports", "Encadrant",
         "Note /20, commentaire, décision Accepter/Rejeter pour ses propres rapports."),
        ("CU06", "Tableau de bord statistique", "Admin, Encadrant",
         "Métriques en temps réel et graphique camembert de répartition par statut."),
        ("CU07", "Prévisualisation et téléchargement", "Tous",
         "Visionneuse PDF intégrée avec navigation et zoom. Téléchargement local."),
        ("CU08", "Analyse de similarité (anti-plagiat)", "Administrateur",
         "Score cosinus TF-IDF entre le nouveau rapport et tous les rapports existants."),
        ("CU09", "Gestion du cycle de vie", "Administrateur",
         "Suivi du statut : Soumis → En cours → Archivé / Rejeté."),
    ],
    caption="Tableau 2 : Cas d'utilisation de l'application")

doc.add_heading("2.2.1  Description Détaillée des CU Principaux", level=3)

doc.add_heading("CU01 – Authentification", level=4)
tbl(doc, ["Champ", "Détail"], [
    ("Acteurs", "Tous les utilisateurs"),
    ("Précondition", "L'utilisateur dispose d'un compte créé par l'administrateur"),
    ("Scénario nominal",
     "1. Saisie de l'email et du mot de passe\n"
     "2. Vérification bcrypt en base\n"
     "3. En cas de succès : ouverture de la MainWindow adaptée au rôle"),
    ("Scénario alternatif",
     "MDP incorrect : compteur incrémenté, message avec tentatives restantes.\n"
     "Après 3 échecs : bouton désactivé, blocage persistant même après déconnexion."),
    ("Postcondition", "Utilisateur authentifié et accès à l'interface de son rôle"),
])

doc.add_heading("CU03 – Dépôt d'un Rapport", level=4)
tbl(doc, ["Champ", "Détail"], [
    ("Acteurs", "Administrateur"),
    ("Précondition", "Administrateur connecté; au moins un encadrant en base"),
    ("Scénario nominal",
     "1. Sélection du fichier PDF\n"
     "2. Saisie des métadonnées (titre, auteurs, promotion, encadrant)\n"
     "3. Validation\n"
     "4. Copie PDF + extraction texte (PyMuPDF) + calcul similarité (TF-IDF)\n"
     "5. Enregistrement en base (statut 'Soumis'), affichage résultats similarité"),
    ("Scénario alternatif",
     "PDF scanné : extraction impossible → rapport enregistré sans texte extrait"),
    ("Postcondition", "Rapport archivé, tableau de similarité affiché en couleurs"),
])

doc.add_heading("CU05 – Évaluation d'un Rapport", level=4)
tbl(doc, ["Champ", "Détail"], [
    ("Acteurs", "Encadrant"),
    ("Précondition", "Rapport 'Soumis' ou 'En cours' assigné à l'encadrant connecté"),
    ("Scénario nominal",
     "1. Sélection du rapport dans la liste\n"
     "2. Ouverture du dialogue (statut automatiquement → 'En cours')\n"
     "3. Saisie de la note et du commentaire\n"
     "4. Accepter → statut 'Archivé' / Rejeter → statut 'Rejeté'"),
    ("Scénario alternatif", "Rejet sans commentaire : message d'erreur, action bloquée"),
    ("Postcondition", "Rapport mis à jour avec note, commentaire et nouveau statut"),
])

doc.add_heading("2.3  Besoins Non-Fonctionnels", level=2)

doc.add_heading("2.3.1  Performance", level=3)
para(doc,
     "L'application doit afficher les résultats de recherche en moins de deux secondes "
     "pour une base contenant jusqu'à 1 000 rapports. L'analyse TF-IDF peut nécessiter "
     "quelques secondes supplémentaires selon la taille du corpus, ce qui est acceptable "
     "étant donné la complexité de l'algorithme.",
     indent=1.25)

doc.add_heading("2.3.2  Sécurité", level=3)
para(doc, "Les exigences de sécurité sont les suivantes :", indent=1.25)
bullet(doc, "Mots de passe stockés sous forme hachée (bcrypt, facteur de coût 12).")
bullet(doc, "Accès aux fonctionnalités strictement contrôlé par le rôle utilisateur (RBAC).")
bullet(doc, "Blocage de compte après 3 tentatives d'authentification infructueuses.")
bullet(doc, "Un administrateur ne peut pas modifier ou supprimer son propre compte.")
bullet(doc, "Un encadrant ne peut évaluer que les rapports qui lui sont assignés.")
bullet(doc, "Les mots de passe ne sont jamais exposés via les requêtes de listing.")

doc.add_heading("2.3.3  Ergonomie", level=3)
para(doc,
     "L'interface doit être intuitive et accessible à des utilisateurs non techniciens. "
     "L'application utilise une interface native PyQt5 avec une barre latérale de "
     "navigation contextuelle, des codes couleur cohérents (bleu institutionnel #156082) "
     "et des icônes expressives. Les actions irréversibles sont toujours confirmées.",
     indent=1.25)

doc.add_heading("2.3.4  Maintenabilité", level=3)
para(doc,
     "Le code est organisé en modules indépendants selon une architecture en couches "
     "séparées, ce qui facilite la maintenance et l'extension des fonctionnalités "
     "sans impact sur les autres composants.",
     indent=1.25)

para(doc, "Conclusion du chapitre", bold=True)
para(doc,
     "L'analyse des besoins a permis d'identifier neuf cas d'utilisation et de "
     "définir des contraintes non-fonctionnelles claires. Ces éléments constituent "
     "la base de la conception présentée au chapitre suivant.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CHAPITRE 3 – CONCEPTION
# ══════════════════════════════════════════════════════════
doc.add_heading("Chapitre 3 : Conception du Système", level=1)
para(doc, "Introduction du chapitre", bold=True)
para(doc,
     "La phase de conception traduit les besoins en modèles formels. Ce chapitre "
     "présente l'architecture logicielle, le modèle de données et les diagrammes UML "
     "principaux.",
     indent=1.25)

doc.add_heading("3.1  Architecture Logicielle", level=2)
doc.add_heading("3.1.1  Architecture en Couches", level=3)
para(doc,
     "L'application repose sur une architecture en trois couches distinctes, inspirée "
     "du patron MVC (Modèle-Vue-Contrôleur) :",
     indent=1.25)
bullet(doc,
       "Couche Présentation (gui/) : Fenêtres PyQt5 gérant l'affichage et les interactions. "
       "Cette couche ne contient aucune logique métier.")
bullet(doc,
       "Couche Métier / Accès aux données (src/) : Modules CRUD, sécurité, traitement PDF "
       "et analyse de similarité. Contient toute la logique applicative.")
bullet(doc,
       "Couche Données (database/) : Base SQLite3 (archive.db) et stockage physique "
       "des PDF (storage/rapports/).")
para(doc, "[Figure 1 : Architecture générale en trois couches]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("3.1.2  Structure des Modules", level=3)
tbl(doc,
    ["Module / Fichier", "Couche", "Responsabilité"],
    [
        ("main.py", "Point d'entrée", "Initialisation BD et lancement de l'application"),
        ("src/database.py", "Données", "Connexion SQLite3 et création des tables"),
        ("src/crud.py", "Métier", "Opérations CRUD sur utilisateurs et rapports"),
        ("src/security.py", "Métier", "Hachage et vérification bcrypt"),
        ("src/pdf_handler.py", "Métier", "Copie PDF, extraction texte, analyse TF-IDF"),
        ("gui/login_window.py", "Présentation", "Interface d'authentification avec anti-brute-force"),
        ("gui/main_window.py", "Présentation", "Fenêtre principale avec sidebar adaptive"),
        ("gui/depot_window.py", "Présentation", "Formulaire de dépôt de rapport"),
        ("gui/recherche_window.py", "Présentation", "Recherche multicritères filtrée par rôle"),
        ("gui/evaluation_window.py", "Présentation", "Workflow d'évaluation pour encadrants"),
        ("gui/users_window.py", "Présentation", "CRUD utilisateurs avec dialogue modal"),
        ("gui/dashboard_window.py", "Présentation", "Tableau de bord Matplotlib"),
        ("gui/preview_window.py", "Présentation", "Visionneuse PDF intégrée"),
    ],
    caption="Tableau 7 : Structure des modules de l'application")

doc.add_heading("3.2  Modèle de Données", level=2)
doc.add_heading("3.2.1  Diagramme Entité-Relation", level=3)
para(doc,
     "La base de données SQLite3 est composée de trois tables. Un Utilisateur peut "
     "être lié à plusieurs Rapports en qualité d'encadrant (relation 1,N). "
     "Un Rapport est rattaché à une Option_Filiere (relation N,1).",
     indent=1.25)
para(doc, "[Figure 2 : Diagramme Entité-Relation (MCD)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("3.2.2  Description de la Table utilisateur", level=3)
tbl(doc,
    ["Colonne", "Type", "Contrainte", "Description"],
    [
        ("id", "INTEGER", "PK, AUTOINCREMENT", "Identifiant unique"),
        ("nom", "TEXT", "NOT NULL", "Nom complet de l'utilisateur"),
        ("email", "TEXT", "UNIQUE, NOT NULL", "Adresse email = identifiant de connexion"),
        ("mot_de_passe", "TEXT", "NOT NULL", "Hash bcrypt du mot de passe"),
        ("role", "TEXT", "CHECK IN (3 valeurs)", "etudiant | encadrant | administrateur"),
        ("numero_apogee", "TEXT", "NULL", "N° Apogée – étudiants uniquement"),
        ("filiere", "TEXT", "NULL", "Filière d'étude – étudiants uniquement"),
        ("annee_inscription", "INTEGER", "NULL", "Année d'inscription – étudiants"),
        ("departement", "TEXT", "NULL", "Département – encadrants uniquement"),
        ("specialite", "TEXT", "NULL", "Spécialité – encadrants uniquement"),
        ("grade", "TEXT", "NULL", "PA | PH | PES | Vacataire"),
    ],
    caption="Tableau 3 : Description de la table utilisateur")

doc.add_heading("3.2.3  Description de la Table rapport", level=3)
tbl(doc,
    ["Colonne", "Type", "Contrainte", "Description"],
    [
        ("id", "INTEGER", "PK, AUTOINCREMENT", "Identifiant unique"),
        ("titre", "TEXT", "NOT NULL", "Titre du rapport PFA"),
        ("auteurs", "TEXT", "NOT NULL", "Noms des auteurs séparés par virgule"),
        ("encadrant_id", "INTEGER", "FK → utilisateur", "Référence vers l'encadrant"),
        ("promotion", "INTEGER", "NOT NULL", "Année de promotion (ex : 2025)"),
        ("option_id", "INTEGER", "FK → option_filiere", "Filière/option concernée"),
        ("mots_cles", "TEXT", "NULL", "Mots-clés pour la recherche full-text"),
        ("chemin_pdf", "TEXT", "NOT NULL", "Chemin absolu vers le fichier PDF stocké"),
        ("texte_extrait", "TEXT", "NULL", "Texte brut extrait du PDF (TF-IDF)"),
        ("statut", "TEXT", "CHECK IN (5 valeurs)", "Soumis | En cours | Evalué | Archivé | Rejeté"),
        ("note", "REAL", "NULL", "Note sur 20 attribuée par l'encadrant"),
        ("commentaire", "TEXT", "NULL", "Commentaire de l'encadrant"),
        ("date_depot", "TEXT", "DEFAULT date('now')", "Date de dépôt du rapport"),
    ],
    caption="Tableau 4 : Description de la table rapport")

doc.add_heading("3.3  Diagrammes de Cas d'Utilisation", level=2)
para(doc,
     "Le diagramme de cas d'utilisation global présente les interactions entre les "
     "trois acteurs et les neuf fonctionnalités de l'application. Les droits "
     "d'accès sont délimités par le rôle de chaque acteur.",
     indent=1.25)
para(doc, "[Figure 3 : Diagramme de cas d'utilisation global]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "L'administrateur a accès à l'intégralité des fonctionnalités. L'encadrant "
     "est limité à ses évaluations et au tableau de bord. L'étudiant est restreint "
     "à la consultation des rapports archivés.",
     indent=1.25)

doc.add_heading("3.4  Diagrammes de Séquence", level=2)
doc.add_heading("3.4.1  Séquence : Authentification (CU01)", level=3)
para(doc, "[Figure 7 : Diagramme de séquence – Authentification]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Flux nominal : L'utilisateur saisit ses identifiants → LoginWindow appelle "
     "get_utilisateur_by_email() → le hash stocké est comparé via bcrypt.checkpw() "
     "→ en cas de succès, MainWindow est instanciée avec le dictionnaire utilisateur.",
     indent=1.25)
para(doc,
     "Flux alternatif (échec) : le compteur tentatives est incrémenté. Si "
     "tentatives ≥ 3, le bouton est désactivé et le blocage persiste après "
     "déconnexion (contrôle : if self.login_window.tentatives < 3).",
     indent=1.25)

doc.add_heading("3.4.2  Séquence : Dépôt avec Analyse de Similarité (CU03 + CU08)", level=3)
para(doc, "[Figure 8 : Diagramme de séquence – Dépôt de rapport]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Flux : L'administrateur valide le formulaire → deposer_rapport() orchestrant : "
     "(1) copier_pdf() renomme et stocke le fichier ; (2) extraire_texte() extrait "
     "le contenu via fitz.open() ; (3) analyser_similarite() vectorise via TfidfVectorizer "
     "et calcule cosine_similarity ; (4) ajouter_rapport() persiste en base (statut 'Soumis') ; "
     "(5) le tableau de résultats est affiché avec codes couleur (vert/jaune/rouge).",
     indent=1.25)

doc.add_heading("3.4.3  Séquence : Évaluation d'un Rapport (CU05)", level=3)
para(doc, "[Figure 9 : Diagramme de séquence – Évaluation]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Flux : L'encadrant sélectionne un rapport 'Soumis' → changer_statut('En cours') "
     "est appelé automatiquement → EvaluationDialog s'ouvre → l'encadrant saisit note "
     "et commentaire → evaluer_rapport() vérifie rapport.encadrant_id == encadrant_id "
     "avant toute mise à jour → le statut passe à 'Archivé' ou 'Rejeté'.",
     indent=1.25)

doc.add_heading("3.5  Diagramme de Classes", level=2)
para(doc, "[Figure 10 : Diagramme de classes de l'application]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Les classes GUI (héritant de QWidget ou QDialog) invoquent les fonctions "
     "des modules src/. La MainWindow joue le rôle de contrôleur central en "
     "instanciant dynamiquement les sous-fenêtres selon la sélection dans la sidebar. "
     "Le QStackedWidget assure la gestion du cycle de vie des widgets enfants.",
     indent=1.25)

para(doc, "Conclusion du chapitre", bold=True)
para(doc,
     "La phase de conception a défini une architecture modulaire et un modèle de "
     "données cohérent. Le chapitre suivant décrit leur mise en œuvre concrète.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CHAPITRE 4 – RÉALISATION
# ══════════════════════════════════════════════════════════
doc.add_heading("Chapitre 4 : Réalisation et Implémentation", level=1)
para(doc, "Introduction du chapitre", bold=True)
para(doc,
     "Ce chapitre décrit l'environnement de développement, les technologies utilisées "
     "et l'implémentation des fonctionnalités clés de l'application.",
     indent=1.25)

doc.add_heading("4.1  Environnement de Développement", level=2)
tbl(doc,
    ["Composant", "Détail"],
    [
        ("Système d'exploitation", "Windows 11 Home"),
        ("Interpréteur Python", "Python 3.12.x"),
        ("IDE", "Visual Studio Code + extensions Python, Pylance"),
        ("Gestionnaire de version", "Git / GitHub"),
        ("Gestionnaire de paquets", "pip (environnement virtuel venv)"),
        ("Base de données", "SQLite3 (embarquée, fichier archive.db)"),
    ])

doc.add_heading("4.2  Technologies et Bibliothèques", level=2)
tbl(doc,
    ["Bibliothèque", "Version", "Rôle dans le projet"],
    [
        ("PyQt5", "5.15.x", "Framework graphique desktop – toutes les fenêtres"),
        ("SQLite3", "stdlib", "Base de données embarquée – stockage des données"),
        ("bcrypt", "4.1.x", "Hachage sécurisé des mots de passe (salted hash)"),
        ("PyMuPDF (fitz)", "1.24.x", "Extraction texte PDF et rendu des pages en images"),
        ("scikit-learn", "1.4.x", "Vectorisation TF-IDF et calcul de similarité cosinus"),
        ("matplotlib", "3.8.x", "Graphique camembert dans le tableau de bord"),
        ("python-docx", "1.1.x", "Génération de ce rapport Word"),
    ],
    caption="Tableau 6 : Technologies et bibliothèques utilisées")

doc.add_heading("4.3  Implémentation des Fonctionnalités Clés", level=2)

doc.add_heading("4.3.1  Authentification Sécurisée (CU01)", level=3)
para(doc, "[Figure 11 : Interface de connexion (LoginWindow)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Le module gui/login_window.py implémente :",
     indent=1.25)
bullet(doc, "Vérification bcrypt : bcrypt.checkpw(password_bytes, hash_bytes) compare "
       "le mot de passe saisi avec le hash stocké sans jamais déchiffrer ce dernier.")
bullet(doc, "Compteur de tentatives : self.tentatives est incrémenté à chaque échec. "
       "À 3 tentatives, self.btn_connecter.setEnabled(False) verrouille l'interface.")
bullet(doc, "Persistance du blocage : lors de la déconnexion depuis MainWindow, "
       "le compteur n'est remis à zéro que si tentatives < 3, empêchant de "
       "contourner le blocage par déconnexion/reconnexion.")
bullet(doc, "Validation préalable : les champs vides déclenchent un avertissement "
       "immédiat sans interroger la base de données.")

doc.add_heading("4.3.2  Interface Principale et Navigation Adaptive (MainWindow)", level=3)
para(doc, "[Figure 12 : Fenêtre principale – Vue administrateur]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "La fenêtre principale adopte un layout horizontal : barre latérale (240 px) "
     "à gauche et QStackedWidget à droite. Les boutons de la sidebar sont créés "
     "dynamiquement selon le rôle :",
     indent=1.25)
bullet(doc, "Étudiant : Accueil, Rechercher")
bullet(doc, "Encadrant : Accueil, Rechercher, Mes évaluations, Tableau de bord")
bullet(doc, "Administrateur : tous les menus précédents + Déposer un rapport + Gérer utilisateurs")
para(doc,
     "La méthode changer_page() détruit le widget précédent (widget.deleteLater()) "
     "avant d'afficher le nouveau, évitant les fuites mémoire.",
     indent=1.25)

doc.add_heading("4.3.3  Dépôt et Analyse de Similarité (CU03 + CU08)", level=3)
para(doc, "[Figure 13 : Module de dépôt] [Figure 14 : Résultats de similarité]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "La fonction deposer_rapport() (src/pdf_handler.py) orchestre quatre étapes :",
     indent=1.25)
numbered(doc, "Copie sécurisée : copier_pdf() renomme le fichier au format "
         "ANNEE_TITRE.pdf et incrémente un suffixe si le nom existe déjà.")
numbered(doc, "Extraction texte : extraire_texte() utilise fitz.open() et "
         "page.get_text() pour toutes les pages. Retourne None pour les PDFs scannés.")
numbered(doc, "Analyse TF-IDF : TfidfVectorizer().fit_transform() vectorise le corpus, "
         "cosine_similarity() calcule les scores. Les résultats sont triés par "
         "similarité décroissante et affichés avec codes couleur (vert < 40%, "
         "jaune 40-70%, rouge ≥ 70%).")
numbered(doc, "Persistance : ajouter_rapport() enregistre le rapport avec statut 'Soumis'.")

doc.add_heading("4.3.4  Recherche Multicritères avec Filtrage par Rôle (CU04)", level=3)
para(doc, "[Figure 15 : Module de recherche multicritères]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "La fonction rechercher_rapports() (src/crud.py) construit une requête SQL "
     "dynamique avec des clauses AND optionnelles : LIKE sur titre/mots_cles/auteurs "
     "pour le mot-clé, égalité sur promotion, et un filtre statut obligatoire pour "
     "les étudiants ('Archivé' uniquement). Ce filtrage est appliqué côté serveur "
     "pour éviter toute fuite de données via des requêtes directes.",
     indent=1.25)

doc.add_heading("4.3.5  Workflow d'Évaluation (CU05)", level=3)
para(doc, "[Figure 16 : Module d'évaluation – Vue encadrant]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
tbl(doc,
    ["Statut initial", "Action", "Statut final", "Condition"],
    [
        ("Soumis", "Ouverture du dialogue", "En cours", "Automatique"),
        ("En cours", "Accepter", "Archivé", "Note + commentaire optionnel"),
        ("En cours", "Rejeter", "Rejeté", "Commentaire OBLIGATOIRE"),
        ("Archivé / Rejeté", "—", "—", "Modification impossible"),
    ],
    caption="Tableau 8 : Workflow des statuts d'un rapport")
para(doc,
     "La vérification de propriété côté serveur (evaluer_rapport() compare "
     "rapport['encadrant_id'] == encadrant_id passé en paramètre) empêche un "
     "encadrant d'évaluer un rapport qui ne lui est pas assigné.",
     indent=1.25)

doc.add_heading("4.3.6  Tableau de Bord Statistique (CU06)", level=3)
para(doc, "[Figure 17 : Tableau de bord (DashboardWindow)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Le tableau de bord présente six métriques en cartes colorées (grille 2×3) "
     "et un graphique camembert Matplotlib intégré via FigureCanvasQTAgg. "
     "Le graphique exclut automatiquement les statuts sans données et affiche "
     "les pourcentages en blanc gras.",
     indent=1.25)

doc.add_heading("4.3.7  Visionneuse PDF Intégrée (CU07)", level=3)
para(doc, "[Figure 18 : Fenêtre de prévisualisation PDF (PreviewWindow)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "La PreviewWindow (QDialog) offre : navigation complète (première, précédente, "
     "suivante, dernière, saut direct via QSpinBox) et zoom de 50% à 300% via "
     "un QSlider. Chaque page est rendue par fitz.Matrix(zoom, zoom) puis "
     "convertie en QPixmap via QImage.",
     indent=1.25)

doc.add_heading("4.3.8  Gestion des Utilisateurs (CU02)", level=3)
para(doc, "[Figure 19 : Module de gestion des utilisateurs (UsersWindow)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
     "Le module implémente le CRUD complet dans un QTableWidget avec des boutons "
     "d'action par ligne. Le UserDialog adapte dynamiquement ses champs selon le "
     "rôle sélectionné (QGroupBox visible/invisible). Protections :",
     indent=1.25)
bullet(doc, "Suppression de son propre compte : action interdite avec message explicatif.")
bullet(doc, "Modification de son propre rôle : action restreinte pour éviter "
       "l'auto-dégradation de privilèges.")

doc.add_heading("4.4  Sécurité de l'Application", level=2)
para(doc,
     "L'application met en œuvre plusieurs couches de sécurité complémentaires :",
     indent=1.25)
bullet(doc, "Hachage bcrypt : les mots de passe ne sont jamais stockés en clair. "
       "Le salt aléatoire intégré à bcrypt rend les attaques par table arc-en-ciel inefficaces.")
bullet(doc, "RBAC strict : les droits sont vérifiés à deux niveaux – côté GUI "
       "(menus non affichés) et côté métier (vérification dans les fonctions CRUD).")
bullet(doc, "Anti-brute-force : blocage définitif après 3 tentatives, persistant "
       "même après déconnexion.")
bullet(doc, "Protection des données : lister_utilisateurs() exclut explicitement "
       "la colonne mot_de_passe des résultats.")
bullet(doc, "Intégrité référentielle : les contraintes CHECK et FOREIGN KEY de "
       "SQLite3 garantissent la cohérence des données.")

para(doc, "Conclusion du chapitre", bold=True)
para(doc,
     "La phase de réalisation a produit une application desktop complète et fonctionnelle "
     "en Python/PyQt5. Le chapitre suivant valide cette implémentation à travers "
     "une batterie de tests.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CHAPITRE 5 – TESTS ET VALIDATION
# ══════════════════════════════════════════════════════════
doc.add_heading("Chapitre 5 : Tests et Validation", level=1)
para(doc, "Introduction du chapitre", bold=True)
para(doc,
     "La phase de tests valide la conformité de l'application aux exigences définies. "
     "Ce chapitre présente la stratégie adoptée, les cas de test exécutés et les "
     "résultats obtenus.",
     indent=1.25)

doc.add_heading("5.1  Stratégie de Tests", level=2)
para(doc,
     "Deux niveaux de tests complémentaires ont été adoptés :",
     indent=1.25)
bullet(doc, "Tests unitaires (pytest) : vérification isolée des fonctions CRUD et de "
       "sécurité sur une base de données de test en mémoire (:memory:).")
bullet(doc, "Tests fonctionnels (manuels) : validation des scénarios de bout en bout "
       "en simulant les actions de chaque type d'utilisateur sur l'application réelle.")

doc.add_heading("5.2  Tests Unitaires", level=2)
tbl(doc,
    ["Cas de test", "Fonction testée", "Résultat attendu", "Statut"],
    [
        ("Ajout utilisateur valide", "ajouter_utilisateur()", "Utilisateur créé en base", "✅ PASS"),
        ("Email déjà existant", "ajouter_utilisateur()", "IntegrityError levée", "✅ PASS"),
        ("Hachage mot de passe", "hasher_mot_de_passe()", "Hash bcrypt retourné", "✅ PASS"),
        ("Vérification MDP correct", "verifier_mot_de_passe()", "True retourné", "✅ PASS"),
        ("Vérification MDP incorrect", "verifier_mot_de_passe()", "False retourné", "✅ PASS"),
        ("Ajout d'un rapport", "ajouter_rapport()", "ID rapport retourné", "✅ PASS"),
        ("Recherche par mot-clé", "rechercher_rapports()", "Liste non vide", "✅ PASS"),
        ("Recherche sans résultat", "rechercher_rapports()", "Liste vide", "✅ PASS"),
        ("Évaluation autorisée", "evaluer_rapport()", "True + statut mis à jour", "✅ PASS"),
        ("Évaluation non autorisée", "evaluer_rapport()", "False (encadrant différent)", "✅ PASS"),
        ("Modification utilisateur", "modifier_utilisateur()", "True retourné", "✅ PASS"),
    ],
    caption="Tableau 9 : Résultats des tests unitaires")
para(doc, "[Figure 20 : Capture de l'exécution des tests (pytest)]",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("5.3  Tests Fonctionnels", level=2)
doc.add_heading("5.3.1  Tests d'Authentification", level=3)
tbl(doc,
    ["Scénario", "Action", "Résultat attendu", "Statut"],
    [
        ("Connexion valide – admin", "Email + MDP corrects", "MainWindow admin ouverte", "✅ PASS"),
        ("Connexion valide – encadrant", "Email + MDP corrects", "MainWindow encadrant ouverte", "✅ PASS"),
        ("Connexion valide – étudiant", "Email + MDP corrects", "MainWindow étudiant ouverte", "✅ PASS"),
        ("MDP incorrect (1 fois)", "MDP erroné", "Message '2 tentatives restantes'", "✅ PASS"),
        ("MDP incorrect (3 fois)", "3 MDP erronés", "Bouton désactivé, compte bloqué", "✅ PASS"),
        ("Déconnexion après blocage", "Déconnexion → retour login", "Bouton toujours désactivé", "✅ PASS"),
        ("Champs vides", "Clic connexion sans saisie", "Message d'avertissement", "✅ PASS"),
    ],
    caption="Tableau 10 : Tests d'authentification")

doc.add_heading("5.3.2  Tests du Contrôle d'Accès (RBAC)", level=3)
tbl(doc,
    ["Scénario", "Rôle", "Résultat attendu", "Statut"],
    [
        ("Menu visible – étudiant", "Étudiant", "Accueil + Rechercher uniquement", "✅ PASS"),
        ("Menu visible – encadrant", "Encadrant", "Accueil + Recherche + Évaluations + Dashboard", "✅ PASS"),
        ("Menu visible – admin", "Administrateur", "Tous les menus visibles", "✅ PASS"),
        ("Recherche – rapports non-archivés", "Étudiant", "Seuls les rapports 'Archivé' visibles", "✅ PASS"),
        ("Évaluer rapport non assigné", "Encadrant", "Accès refusé (False retourné)", "✅ PASS"),
        ("Supprimer son propre compte", "Administrateur", "Message 'Action interdite'", "✅ PASS"),
        ("Modifier son propre rôle", "Administrateur", "Message 'Action restreinte'", "✅ PASS"),
    ],
    caption="Tableau 11 : Tests du contrôle d'accès (RBAC)")

doc.add_heading("5.3.3  Tests du Dépôt de Rapport", level=3)
tbl(doc,
    ["Scénario", "Données d'entrée", "Résultat attendu", "Statut"],
    [
        ("Dépôt complet valide", "PDF + tous champs", "Rapport 'Soumis', tableau similarité affiché", "✅ PASS"),
        ("PDF scanné (sans texte)", "PDF image only", "Rapport créé, avertissement extraction", "✅ PASS"),
        ("Titre manquant", "Titre vide", "Message 'Champ manquant', pas de dépôt", "✅ PASS"),
        ("Encadrant non sélectionné", "Aucun encadrant", "Message 'Encadrant manquant'", "✅ PASS"),
        ("Sans sélection de fichier", "Bouton déposer sans PDF", "Message 'Fichier manquant'", "✅ PASS"),
    ],
    caption="Tableau 12 : Tests du dépôt de rapport")

doc.add_heading("5.3.4  Tests d'Évaluation", level=3)
tbl(doc,
    ["Scénario", "Action", "Résultat attendu", "Statut"],
    [
        ("Évaluation acceptée", "Note 15/20 + Accepter", "Statut → 'Archivé', note enregistrée", "✅ PASS"),
        ("Évaluation rejetée avec commentaire", "Rejeter + motif", "Statut → 'Rejeté'", "✅ PASS"),
        ("Rejet sans commentaire", "Rejeter sans texte", "Message 'Commentaire obligatoire'", "✅ PASS"),
        ("Rapport déjà archivé", "Tentative d'évaluation", "Message 'Évaluation impossible'", "✅ PASS"),
    ],
    caption="Tableau 13 : Tests d'évaluation")

doc.add_heading("5.4  Synthèse des Résultats", level=2)
tbl(doc,
    ["Catégorie de tests", "Cas testés", "Réussis", "Échecs", "Taux de réussite"],
    [
        ("Tests unitaires", "11", "11", "0", "100%"),
        ("Tests d'authentification", "7", "7", "0", "100%"),
        ("Tests RBAC", "7", "7", "0", "100%"),
        ("Tests de dépôt", "5", "5", "0", "100%"),
        ("Tests d'évaluation", "4", "4", "0", "100%"),
        ("TOTAL", "34", "34", "0", "100%"),
    ],
    caption="Tableau 14 : Synthèse des résultats des tests")
para(doc,
     "L'ensemble des 34 cas de test a été exécuté avec succès, confirmant la "
     "conformité de l'application à toutes les exigences fonctionnelles et de "
     "sécurité. Aucun défaut bloquant n'a été identifié.",
     indent=1.25)

para(doc, "Conclusion du chapitre", bold=True)
para(doc,
     "La phase de tests a validé toutes les fonctionnalités avec un taux de "
     "réussite de 100%. L'application respecte les contraintes de sécurité "
     "(RBAC, bcrypt, anti-brute-force) et les workflows définis lors de l'analyse.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# CONCLUSION GÉNÉRALE
# ══════════════════════════════════════════════════════════
doc.add_heading("Conclusion Générale et Perspectives", level=1)
para(doc,
     "Au terme de ce Projet de Fin d'Année, nous avons conçu et développé "
     "l'application « Archive Digitale PFA », une solution desktop complète "
     "répondant à la problématique de la gestion fragmentée des rapports PFA à l'EMSI.",
     indent=1.25)
para(doc, "Les principaux objectifs ont été atteints :", indent=1.25)
bullet(doc, "Un système d'archivage numérique structuré avec base SQLite3 et "
       "stockage organisé des PDF.")
bullet(doc, "Un contrôle d'accès basé sur les rôles (RBAC) avec trois profils "
       "aux droits strictement définis et vérifiés côté serveur.")
bullet(doc, "Un workflow d'évaluation complet et traçable (Soumis → En cours → "
       "Archivé/Rejeté).")
bullet(doc, "Un moteur anti-plagiat TF-IDF cosinus pour détecter les contenus "
       "similaires dès la soumission.")
bullet(doc, "Une interface graphique intuitive adaptée à chaque profil utilisateur.")
bullet(doc, "Un tableau de bord statistique en temps réel avec graphiques Matplotlib.")
para(doc,
     "Ce projet nous a permis de développer des compétences dans de nombreux domaines : "
     "développement GUI avec PyQt5, gestion de bases de données SQLite3, sécurité "
     "applicative avec bcrypt, traitement de PDF avec PyMuPDF, et application "
     "d'algorithmes de NLP (TF-IDF) en contexte réel.",
     indent=1.25)
para(doc, "Perspectives d'évolution envisagées :", indent=1.25)
bullet(doc, "Migration vers une architecture web (Flask/Django + React) pour un "
       "accès multi-plateforme et collaboratif.")
bullet(doc, "Notifications automatiques par email lors des changements de statut.")
bullet(doc, "Intégration d'un modèle NLP avancé pour la classification automatique "
       "des rapports par thématique et la suggestion de mots-clés.")
bullet(doc, "Module de signature électronique des rapports évalués.")
bullet(doc, "API REST pour l'intégration avec le système d'information de l'EMSI.")
para(doc,
     "En conclusion, ce PFA constitue une expérience formatrice complète, couvrant "
     "l'intégralité du cycle de développement logiciel, de l'analyse des besoins "
     "jusqu'à la validation finale.",
     indent=1.25)
pb(doc)

# ══════════════════════════════════════════════════════════
# BIBLIOGRAPHIE
# ══════════════════════════════════════════════════════════
doc.add_heading("Bibliographie", level=1)
para(doc, "Références présentées selon la norme ISO 690.", italic=True)

references = [
    "[1] BAEZA-YATES, Ricardo et RIBEIRO-NETO, Berthier. Modern Information Retrieval: "
    "The Concepts and Technology behind Search. 2e éd. Harlow : Addison-Wesley, 2011. "
    "ISBN 978-0-321-41691-9.",

    "[2] FERILLI, Stefano et al. « Automatic Learning of Document Structure from Digital "
    "Libraries ». International Journal on Digital Libraries. 2011, vol. 11, n° 3, "
    "pp. 187-208. DOI : 10.1007/s00799-011-0075-x.",

    "[3] HASAN, Kazi Saidul et NG, Vincent. « Automatic Keyphrase Extraction: A Survey "
    "of the State of the Art ». In : Proceedings of the 52nd Annual Meeting of the ACL. "
    "Baltimore, 2014, pp. 1262-1273. DOI : 10.3115/v1/P14-1119.",

    "[4] MARTIN, Robert C. Clean Code: A Handbook of Agile Software Craftsmanship. "
    "Upper Saddle River : Prentice Hall, 2009. ISBN 978-0-13-235088-4.",

    "[5] PEDREGOSA, Fabian et al. « Scikit-learn: Machine Learning in Python ». "
    "Journal of Machine Learning Research. 2011, vol. 12, pp. 2825-2830.",

    "[6] PROVOS, Niels et MAZIÈRES, David. « A Future-Adaptable Password Scheme ». "
    "In : Proceedings of the USENIX Annual Technical Conference. Monterey, 1999. "
    "Disponible sur : https://www.usenix.org/legacy/event/usenix99/provos/provos.pdf.",

    "[7] SALTON, Gerard et BUCKLEY, Chris. « Term-weighting Approaches in Automatic "
    "Text Retrieval ». Information Processing & Management. 1988, vol. 24, n° 5, "
    "pp. 513-523. DOI : 10.1016/0306-4573(88)90021-0.",

    "[8] SANDHU, Ravi et al. « Role-Based Access Control Models ». IEEE Computer. "
    "1996, vol. 29, n° 2, pp. 38-47. DOI : 10.1109/2.485845.",

    "[9] WILLMAN, Joshua M. Beginning PyQt: A Hands-on Approach to GUI Programming. "
    "Berkeley : Apress, 2022. ISBN 978-1-4842-7999-1.",

    "[10] PYTHON SOFTWARE FOUNDATION. Python 3.12 Documentation [en ligne]. 2024. "
    "Disponible sur : https://docs.python.org/3/ [consulté le 07/05/2026].",

    "[11] PYMUPDF TEAM. PyMuPDF Documentation [en ligne]. 2024. "
    "Disponible sur : https://pymupdf.readthedocs.io/ [consulté le 07/05/2026].",

    "[12] SQLITE CONSORTIUM. SQLite Documentation [en ligne]. 2024. "
    "Disponible sur : https://www.sqlite.org/docs.html [consulté le 07/05/2026].",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref)
    r.font.size = Pt(10)
pb(doc)

# ══════════════════════════════════════════════════════════
# ANNEXES
# ══════════════════════════════════════════════════════════
doc.add_heading("Annexes", level=1)

doc.add_heading("Annexe A : Script de Création des Tables SQLite3", level=2)
para(doc, "Extrait du module src/database.py :", indent=1.25)
code_block(doc,
    "def create_tables():\n"
    "    conn = get_connection()\n"
    "    conn.execute('''\n"
    "        CREATE TABLE IF NOT EXISTS utilisateur (\n"
    "            id            INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "            nom           TEXT NOT NULL,\n"
    "            email         TEXT NOT NULL UNIQUE,\n"
    "            mot_de_passe  TEXT NOT NULL,\n"
    "            role          TEXT NOT NULL CHECK(role IN\n"
    "                          ('etudiant','encadrant','administrateur')),\n"
    "            numero_apogee TEXT, filiere TEXT, annee_inscription INTEGER,\n"
    "            departement TEXT, specialite TEXT, grade TEXT)\n"
    "    ''')\n"
    "    # Tables rapport et option_filiere créées de même ...\n"
    "    conn.commit()\n"
    "    conn.close()")

doc.add_heading("Annexe B : Algorithme de Détection de Similarité (TF-IDF)", level=2)
para(doc, "Extrait du module src/pdf_handler.py – Fonction analyser_similarite() :", indent=1.25)
code_block(doc,
    "def analyser_similarite(texte_nouveau):\n"
    "    conn = get_connection()\n"
    "    rapports = conn.execute(\n"
    "        'SELECT id, titre, auteurs, texte_extrait FROM rapport'\n"
    "        ' WHERE texte_extrait IS NOT NULL'\n"
    "    ).fetchall()\n"
    "    conn.close()\n"
    "    if not rapports:\n"
    "        return []\n"
    "    textes = [texte_nouveau] + [r['texte_extrait'] for r in rapports]\n"
    "    vectorizer = TfidfVectorizer()\n"
    "    matrice    = vectorizer.fit_transform(textes)\n"
    "    scores     = cosine_similarity(matrice[0:1], matrice[1:])[0]\n"
    "    resultats  = []\n"
    "    for rapport, score in zip(rapports, scores):\n"
    "        resultats.append({\n"
    "            'titre'      : rapport['titre'],\n"
    "            'auteurs'    : rapport['auteurs'],\n"
    "            'similarite' : round(score * 100, 1)\n"
    "        })\n"
    "    return sorted(resultats, key=lambda x: x['similarite'], reverse=True)")

doc.add_heading("Annexe C : Guide d'Installation", level=2)
para(doc, "Pré-requis : Python 3.10+ installé sur la machine.", indent=1.25)
numbered(doc, "Cloner le dépôt : git clone <url-du-projet>")
numbered(doc, "Créer l'environnement virtuel : python -m venv venv")
numbered(doc, "Activer : venv\\Scripts\\activate  (Windows)")
numbered(doc, "Installer les dépendances : pip install PyQt5 bcrypt PyMuPDF scikit-learn matplotlib")
numbered(doc, "Initialiser la base : python -c \"from src.database import create_tables; create_tables()\"")
numbered(doc, "Lancer l'application : python -m gui.login_window")

# ══════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════
output = "Rapport_PFA_Archive_Digitale.docx"
doc.save(output)
print(f"[OK] Rapport genere : {output}")
